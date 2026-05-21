from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

from backend.models import CompositeAssignment, VariantSet, WorkSheet


def _build_students(ws: WorkSheet) -> Document:
    doc = Document()
    doc.add_heading(f"{ws.subject}, {ws.grade} класс. {ws.topic}", level=1)
    for variant in ws.variants:
        doc.add_heading(f"Вариант {variant.number}", level=2)
        for i, task in enumerate(variant.tasks, 1):
            doc.add_paragraph(f"{i}. {task.statement}")
        doc.add_page_break()
    return doc


def _build_teacher(ws: WorkSheet) -> Document:
    doc = Document()
    doc.add_heading(f"{ws.subject}, {ws.grade} класс. {ws.topic} - ключи", level=1)
    for variant in ws.variants:
        doc.add_heading(f"Вариант {variant.number}", level=2)
        for i, task in enumerate(variant.tasks, 1):
            doc.add_paragraph(f"{i}. {task.statement}")
            p = doc.add_paragraph()
            p.add_run(f"Ответ: {task.answer}").italic = True
            doc.add_paragraph(f"Решение: {task.solution}")
            doc.add_paragraph(f"Критерии: {task.grading_criteria}")
        doc.add_page_break()
    return doc


def export_for_students(ws: WorkSheet, out_path: Path) -> Path:
    _build_students(ws).save(out_path)
    return out_path


def export_for_teacher(ws: WorkSheet, out_path: Path) -> Path:
    _build_teacher(ws).save(out_path)
    return out_path


def students_docx_bytes(ws: WorkSheet) -> bytes:
    buf = BytesIO()
    _build_students(ws).save(buf)
    return buf.getvalue()


def teacher_docx_bytes(ws: WorkSheet) -> bytes:
    buf = BytesIO()
    _build_teacher(ws).save(buf)
    return buf.getvalue()


# ---------- Case 4: VariantSet ----------


def _variant_set_title(vs: VariantSet) -> str:
    ref = vs.reference
    parts = [ref.subject]
    if ref.grade:
        parts.append(f"{ref.grade} класс")
    parts.append(ref.topic)
    return ". ".join(parts)


def _build_variants_students(vs: VariantSet) -> Document:
    doc = Document()
    doc.add_heading(_variant_set_title(vs), level=1)
    info = doc.add_paragraph()
    info.add_run(f"Вариантов: {len(vs.variants)}. ").bold = True
    info.add_run("Эталон: ").italic = True
    info.add_run(vs.reference.raw_statement).italic = True

    for v in vs.variants:
        doc.add_heading(f"Вариант {v.number}", level=2)
        meta = doc.add_paragraph()
        meta.add_run(f"Сложность: {v.difficulty}/5").italic = True
        body = doc.add_paragraph()
        body.add_run(v.statement)
        # Место для ФИО ученика и ответа
        line = doc.add_paragraph()
        line.add_run("\nФИО ученика: ____________________________   Класс: _______").italic = True
        ans = doc.add_paragraph()
        ans.add_run("Ответ: ____________________________").italic = True
        doc.add_page_break()
    return doc


def _build_variants_teacher(vs: VariantSet) -> Document:
    doc = Document()
    doc.add_heading(f"{_variant_set_title(vs)} - ключи учителя", level=1)
    info = doc.add_paragraph()
    info.add_run(f"Вариантов: {len(vs.variants)}.").bold = True

    eth = doc.add_paragraph()
    eth.add_run("Эталон: ").bold = True
    eth.add_run(vs.reference.raw_statement)

    tmpl = doc.add_paragraph()
    tmpl.add_run("Шаблон: ").bold = True
    tmpl.add_run(vs.reference.template).font.name = "Consolas"

    fm = doc.add_paragraph()
    fm.add_run("Формула ответа: ").bold = True
    fm.add_run(vs.reference.answer_formula).font.name = "Consolas"

    if vs.reference.slots:
        sp = doc.add_paragraph()
        sp.add_run("Слоты: ").bold = True
        descr = ", ".join(
            (
                f"{s.name} ({s.type.value}"
                + (
                    f", диапазон {int(s.range[0])}…{int(s.range[1])}"
                    if s.range
                    else ""
                )
                + (", заблокирован" if s.locked else "")
                + ")"
            )
            for s in vs.reference.slots
        )
        sp.add_run(descr).font.size = Pt(10)

    doc.add_paragraph()

    for v in vs.variants:
        doc.add_heading(f"Вариант {v.number}", level=2)

        body = doc.add_paragraph()
        body.add_run(v.statement)

        meta = doc.add_paragraph()
        meta.add_run("Ответ: ").bold = True
        meta.add_run(v.answer)
        meta.add_run(f"   ·   Сложность: {v.difficulty}/5")
        meta.add_run(
            f"   ·   Sympy: {'сверено' if v.sympy_verified else 'требуется проверка'}"
        )

        values = doc.add_paragraph()
        values.add_run("Подставленные значения: ").italic = True
        values_str = ", ".join(f"{k}={vv}" for k, vv in v.slot_values.items())
        values.add_run(values_str).italic = True

        if v.solution:
            sol = doc.add_paragraph()
            sol.add_run("Решение: ").bold = True
            sol.add_run(v.solution)

        if v.issues:
            warn = doc.add_paragraph()
            warn.add_run("Замечания: ").bold = True
            warn.add_run("; ".join(v.issues))
    return doc


def variants_docx_students(vs: VariantSet) -> bytes:
    buf = BytesIO()
    _build_variants_students(vs).save(buf)
    return buf.getvalue()


def variants_docx_teacher(vs: VariantSet) -> bytes:
    buf = BytesIO()
    _build_variants_teacher(vs).save(buf)
    return buf.getvalue()


# ---------- Composite (контрольная из нескольких блоков) ----------


def _assignment_title(a: CompositeAssignment) -> str:
    parts = []
    for b in a.blocks:
        parts.append(f"{b.topic} ({b.grade} кл)")
    if len(parts) == 1:
        return f"Контрольная: {parts[0]}"
    return "Контрольная: " + " + ".join(parts)


def _build_assignment_students(a: CompositeAssignment) -> Document:
    doc = Document()
    doc.add_heading(_assignment_title(a), level=1)
    info = doc.add_paragraph()
    info.add_run(f"Вариантов: {a.variants_count}.").bold = True
    info.add_run(
        f" Задач в варианте: {sum(b.tasks_per_variant for b in a.blocks)}."
    )

    for v in range(a.variants_count):
        doc.add_heading(f"Вариант {v + 1}", level=2)
        task_number = 1
        for block in a.blocks:
            sub = doc.add_paragraph()
            sub.add_run(f"{block.topic} ({block.grade} класс)").bold = True
            for task in block.variant_tasks[v]:
                p = doc.add_paragraph()
                p.add_run(f"{task_number}. {task.statement}")
                task_number += 1
        sign = doc.add_paragraph()
        sign.add_run(
            "\nФИО ученика: ____________________________   Класс: _______"
        ).italic = True
        doc.add_page_break()
    return doc


def _build_assignment_teacher(a: CompositeAssignment) -> Document:
    doc = Document()
    doc.add_heading(_assignment_title(a) + " — ключи учителя", level=1)

    info = doc.add_paragraph()
    info.add_run(f"Вариантов: {a.variants_count}.").bold = True
    info.add_run(f" Блоков: {len(a.blocks)}.")

    for i, block in enumerate(a.blocks, 1):
        sec = doc.add_paragraph()
        sec.add_run(f"Блок {i}: {block.topic} (класс {block.grade})").bold = True
        sec.add_run(f"  ·  задач в варианте: {block.tasks_per_variant}").italic = True

        eth = doc.add_paragraph()
        eth.add_run("Эталон: ").bold = True
        eth.add_run(block.reference.raw_statement)

        fm = doc.add_paragraph()
        fm.add_run("Формула: ").bold = True
        fm.add_run(block.reference.answer_formula).font.name = "Consolas"

    doc.add_paragraph()

    for v in range(a.variants_count):
        doc.add_heading(f"Вариант {v + 1}", level=2)
        task_number = 1
        for block in a.blocks:
            sub = doc.add_paragraph()
            sub.add_run(f"  {block.topic} ({block.grade} класс)").italic = True
            for task in block.variant_tasks[v]:
                p = doc.add_paragraph()
                p.add_run(f"{task_number}. {task.statement}")
                task_number += 1

                meta = doc.add_paragraph()
                meta.add_run("    Ответ: ").bold = True
                meta.add_run(task.answer)
                meta.add_run(f"   ·   Сложность: {task.difficulty}/5")
                meta.add_run(
                    f"   ·   Sympy: {'сверено' if task.sympy_verified else 'требуется проверка'}"
                )

                values = doc.add_paragraph()
                values.add_run("    Значения: ").italic = True
                values_str = ", ".join(f"{k}={vv}" for k, vv in task.slot_values.items())
                values.add_run(values_str).italic = True
    return doc


def assignment_docx_students(a: CompositeAssignment) -> bytes:
    buf = BytesIO()
    _build_assignment_students(a).save(buf)
    return buf.getvalue()


def assignment_docx_teacher(a: CompositeAssignment) -> bytes:
    buf = BytesIO()
    _build_assignment_teacher(a).save(buf)
    return buf.getvalue()
